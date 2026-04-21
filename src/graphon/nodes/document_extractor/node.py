from __future__ import annotations

import csv
import io
import json
import logging
import pathlib
import tempfile
import zipfile
from collections.abc import Callable, Mapping, Sequence
from dataclasses import dataclass
from typing import Any, override

import charset_normalizer
import docx
import pandas as pd
import pypandoc
import pypdfium2
import webvtt
import yaml
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from graphon.entities.graph_init_params import GraphInitParams
from graphon.enums import BuiltinNodeTypes, WorkflowNodeExecutionStatus
from graphon.file import file_manager
from graphon.file.enums import FileTransferMethod
from graphon.file.models import File
from graphon.http import HttpClientProtocol, get_http_client
from graphon.node_events.base import NodeRunResult
from graphon.nodes.base.node import Node
from graphon.runtime.graph_runtime_state import GraphRuntimeState
from graphon.variables.segments import ArrayFileSegment, ArrayStringSegment, FileSegment

from .entities import DocumentExtractorNodeData, UnstructuredApiConfig
from .exc import (
    DocumentExtractorError,
    FileDownloadError,
    TextExtractionError,
    UnsupportedFileTypeError,
)

logger = logging.getLogger(__name__)

_MIME_PLAIN_TEXT_TYPES = frozenset((
    "text/plain",
    "text/html",
    "text/htm",
    "text/markdown",
    "text/xml",
))
_EXTENSION_PLAIN_TEXT_TYPES = frozenset((
    ".txt",
    ".markdown",
    ".md",
    ".mdx",
    ".html",
    ".htm",
    ".xml",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
    ".cc",
    ".cxx",
    ".c++",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".java",
    ".php",
    ".rb",
    ".go",
    ".rs",
    ".swift",
    ".kt",
    ".scala",
    ".sh",
    ".bash",
    ".bat",
    ".ps1",
    ".sql",
    ".r",
    ".m",
    ".pl",
    ".lua",
    ".vim",
    ".asm",
    ".s",
    ".css",
    ".scss",
    ".less",
    ".sass",
    ".ini",
    ".cfg",
    ".conf",
    ".toml",
    ".env",
    ".log",
    ".vtt",
))


@dataclass(frozen=True)
class _ExtractorRegistration:
    name: str
    extractor: Callable[..., str]
    mime_types: frozenset[str] = frozenset()
    file_extensions: frozenset[str] = frozenset()
    requires_unstructured_config: bool = False

    def extract(
        self,
        *,
        file_content: bytes,
        unstructured_api_config: UnstructuredApiConfig,
    ) -> str:
        if self.requires_unstructured_config:
            return self.extractor(
                file_content,
                unstructured_api_config=unstructured_api_config,
            )
        return self.extractor(file_content)


class _TextExtractorRegistry:
    def __init__(self, registrations: Sequence[_ExtractorRegistration]) -> None:
        self._mime_type_extractors = self._build_lookup(
            registrations,
            kind_selector=lambda registration: registration.mime_types,
        )
        self._file_extension_extractors = self._build_lookup(
            registrations,
            kind_selector=lambda registration: registration.file_extensions,
        )

    def extract_by_mime_type(
        self,
        *,
        file_content: bytes,
        mime_type: str,
        unstructured_api_config: UnstructuredApiConfig,
    ) -> str:
        return self._extract(
            extractor_lookup=self._mime_type_extractors,
            file_content=file_content,
            file_kind=mime_type,
            unstructured_api_config=unstructured_api_config,
            error_label="Unsupported MIME type",
        )

    def extract_by_file_extension(
        self,
        *,
        file_content: bytes,
        file_extension: str,
        unstructured_api_config: UnstructuredApiConfig,
    ) -> str:
        return self._extract(
            extractor_lookup=self._file_extension_extractors,
            file_content=file_content,
            file_kind=file_extension,
            unstructured_api_config=unstructured_api_config,
            error_label="Unsupported Extension Type",
        )

    @staticmethod
    def _build_lookup(
        registrations: Sequence[_ExtractorRegistration],
        *,
        kind_selector: Callable[[_ExtractorRegistration], frozenset[str]],
    ) -> dict[str, _ExtractorRegistration]:
        lookup: dict[str, _ExtractorRegistration] = {}
        for registration in registrations:
            for kind in kind_selector(registration):
                if kind in lookup:
                    current_registration = lookup[kind]
                    msg = (
                        f"Duplicate extractor registration for {kind!r}: "
                        f"{current_registration.name} and {registration.name}"
                    )
                    raise ValueError(msg)
                lookup[kind] = registration
        return lookup

    @staticmethod
    def _extract(
        *,
        extractor_lookup: Mapping[str, _ExtractorRegistration],
        file_content: bytes,
        file_kind: str,
        unstructured_api_config: UnstructuredApiConfig,
        error_label: str,
    ) -> str:
        registration = extractor_lookup.get(file_kind)
        if registration is None:
            msg = f"{error_label}: {file_kind}"
            raise UnsupportedFileTypeError(msg)
        return registration.extract(
            file_content=file_content,
            unstructured_api_config=unstructured_api_config,
        )


def _partition_file_via_unstructured_api(
    partition_via_api: Any,
    file_content: bytes,
    *,
    suffix: str,
    unstructured_api_config: UnstructuredApiConfig,
) -> Sequence[Any]:
    api_key = unstructured_api_config.api_key or ""

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
        temp_file.write(file_content)
        temp_file.flush()
        temp_path = pathlib.Path(temp_file.name)

    try:
        with temp_path.open("rb") as file:
            return partition_via_api(
                file=file,
                metadata_filename=temp_path.name,
                api_url=unstructured_api_config.api_url,
                api_key=api_key,
            )
    finally:
        temp_path.unlink(missing_ok=True)


class DocumentExtractorNode(Node[DocumentExtractorNodeData]):
    """Extracts text content from various file types.
    Supports plain text, PDF, and DOC/DOCX files.
    """

    node_type = BuiltinNodeTypes.DOCUMENT_EXTRACTOR

    @classmethod
    @override
    def version(cls) -> str:
        return "1"

    @override
    def __init__(
        self,
        node_id: str,
        data: DocumentExtractorNodeData,
        *,
        graph_init_params: GraphInitParams,
        graph_runtime_state: GraphRuntimeState,
        unstructured_api_config: UnstructuredApiConfig | None = None,
        http_client: HttpClientProtocol | None = None,
    ) -> None:
        super().__init__(
            node_id=node_id,
            data=data,
            graph_init_params=graph_init_params,
            graph_runtime_state=graph_runtime_state,
        )
        self._unstructured_api_config = (
            unstructured_api_config or UnstructuredApiConfig()
        )
        self._http_client = http_client or get_http_client()

    @property
    def http_client(self) -> HttpClientProtocol:
        """Return the HTTP client used to retrieve file contents."""
        return self._http_client

    @override
    def _run(self) -> NodeRunResult:
        variable_selector = self.node_data.variable_selector
        variable = self.graph_runtime_state.variable_pool.get(variable_selector)

        error_message = None
        if variable is None:
            error_message = f"File variable not found for selector: {variable_selector}"
        elif variable.value and not isinstance(
            variable,
            ArrayFileSegment | FileSegment,
        ):
            error_message = f"Variable {variable_selector} is not an ArrayFileSegment"

        if error_message is not None:
            return NodeRunResult(
                status=WorkflowNodeExecutionStatus.FAILED,
                error=error_message,
            )

        if variable is None:
            msg = f"File variable not found for selector: {variable_selector}"
            return NodeRunResult(
                status=WorkflowNodeExecutionStatus.FAILED,
                error=msg,
            )
        value = variable.value
        inputs = {"variable_selector": variable_selector}
        if isinstance(value, list):
            value = list(filter(lambda x: x, value))
        process_data = {"documents": value if isinstance(value, list) else [value]}

        if not value:
            return NodeRunResult(
                status=WorkflowNodeExecutionStatus.SUCCEEDED,
                inputs=inputs,
                process_data=process_data,
                outputs={"text": ArrayStringSegment(value=[])},
            )

        if isinstance(value, list):
            try:
                extracted_text_list = [
                    _extract_text_from_file(
                        self._http_client,
                        file,
                        unstructured_api_config=self._unstructured_api_config,
                    )
                    for file in value
                ]
            except DocumentExtractorError as e:
                logger.warning(e, exc_info=True)
                return NodeRunResult(
                    status=WorkflowNodeExecutionStatus.FAILED,
                    error=str(e),
                    inputs=inputs,
                    process_data=process_data,
                )
            outputs = {"text": ArrayStringSegment(value=extracted_text_list)}
        else:
            try:
                extracted_text = _extract_text_from_file(
                    self._http_client,
                    value,
                    unstructured_api_config=self._unstructured_api_config,
                )
            except DocumentExtractorError as e:
                logger.warning(e, exc_info=True)
                return NodeRunResult(
                    status=WorkflowNodeExecutionStatus.FAILED,
                    error=str(e),
                    inputs=inputs,
                    process_data=process_data,
                )
            outputs = {"text": extracted_text}

        return NodeRunResult(
            status=WorkflowNodeExecutionStatus.SUCCEEDED,
            inputs=inputs,
            process_data=process_data,
            outputs=outputs,
        )

    @classmethod
    @override
    def _extract_variable_selector_to_variable_mapping(
        cls,
        *,
        graph_config: Mapping[str, Any],
        node_id: str,
        node_data: DocumentExtractorNodeData,
    ) -> Mapping[str, Sequence[str]]:
        _ = graph_config  # Explicitly mark as unused
        return {node_id + ".files": node_data.variable_selector}


def _extract_text_by_mime_type(
    *,
    file_content: bytes,
    mime_type: str,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    """Extract text from a file based on its MIME type."""
    return _TEXT_EXTRACTOR_REGISTRY.extract_by_mime_type(
        file_content=file_content,
        mime_type=mime_type,
        unstructured_api_config=unstructured_api_config,
    )


def _extract_text_by_file_extension(
    *,
    file_content: bytes,
    file_extension: str,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    """Extract text from a file based on its file extension."""
    return _TEXT_EXTRACTOR_REGISTRY.extract_by_file_extension(
        file_content=file_content,
        file_extension=file_extension,
        unstructured_api_config=unstructured_api_config,
    )


def _extract_text_from_plain_text(file_content: bytes) -> str:
    try:
        # Detect encoding using charset_normalizer
        result = charset_normalizer.from_bytes(
            file_content,
            cp_isolation=["utf_8", "latin_1", "cp1252"],
        ).best()
        encoding = result.encoding if result else "utf-8"

        # Fallback to utf-8 if detection fails
        if not encoding:
            encoding = "utf-8"

        return file_content.decode(encoding, errors="ignore")
    except (UnicodeDecodeError, LookupError) as e:
        # If decoding fails, try with utf-8 as last resort
        try:
            return file_content.decode("utf-8", errors="ignore")
        except UnicodeDecodeError:
            msg = f"Failed to decode plain text file: {e}"
            raise TextExtractionError(msg) from e


def _extract_text_from_json(file_content: bytes) -> str:
    try:
        # Detect encoding using charset_normalizer
        result = charset_normalizer.from_bytes(file_content).best()
        encoding = result.encoding if result else "utf-8"

        # Fallback to utf-8 if detection fails
        if not encoding:
            encoding = "utf-8"

        json_data = json.loads(file_content.decode(encoding, errors="ignore"))
        return json.dumps(json_data, indent=2, ensure_ascii=False)
    except (UnicodeDecodeError, LookupError, json.JSONDecodeError) as e:
        # If decoding fails, try with utf-8 as last resort
        try:
            json_data = json.loads(file_content.decode("utf-8", errors="ignore"))
            return json.dumps(json_data, indent=2, ensure_ascii=False)
        except (UnicodeDecodeError, json.JSONDecodeError):
            msg = f"Failed to decode or parse JSON file: {e}"
            raise TextExtractionError(msg) from e


def _extract_text_from_yaml(file_content: bytes) -> str:
    """Extract the content from yaml file"""
    try:
        # Detect encoding using charset_normalizer
        result = charset_normalizer.from_bytes(file_content).best()
        encoding = result.encoding if result else "utf-8"

        # Fallback to utf-8 if detection fails
        if not encoding:
            encoding = "utf-8"

        yaml_data = yaml.safe_load_all(file_content.decode(encoding, errors="ignore"))
        return yaml.dump_all(yaml_data, allow_unicode=True, sort_keys=False)
    except (UnicodeDecodeError, LookupError, yaml.YAMLError) as e:
        # If decoding fails, try with utf-8 as last resort
        try:
            yaml_data = yaml.safe_load_all(
                file_content.decode("utf-8", errors="ignore"),
            )
            return yaml.dump_all(yaml_data, allow_unicode=True, sort_keys=False)
        except (UnicodeDecodeError, yaml.YAMLError):
            msg = f"Failed to decode or parse YAML file: {e}"
            raise TextExtractionError(msg) from e


def _extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_document = pypdfium2.PdfDocument(pdf_file, autoclose=True)
        text = ""
        for page in pdf_document:
            text_page = page.get_textpage()
            text += text_page.get_text_range()
            text_page.close()
            page.close()
    except Exception as e:
        msg = f"Failed to extract text from PDF: {e!s}"
        raise TextExtractionError(msg) from e
    else:
        return text


def _extract_text_from_doc(
    file_content: bytes,
    *,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    """Extract text from a DOC file."""
    from unstructured.partition.api import partition_via_api  # noqa: PLC0415

    if not unstructured_api_config.api_url:
        msg = "Unstructured API URL is not configured for DOC file processing."
        raise TextExtractionError(msg)

    try:
        elements = _partition_file_via_unstructured_api(
            partition_via_api,
            file_content,
            suffix=".doc",
            unstructured_api_config=unstructured_api_config,
        )
        return "\n".join([getattr(element, "text", "") for element in elements])
    except Exception as e:
        msg = f"Failed to extract text from DOC: {e!s}"
        raise TextExtractionError(msg) from e


def parser_docx_part(
    block: object,
    doc: Document,
    content_items: list[tuple[int, str, Table | Paragraph]],
    i: int,
) -> None:
    content_item: tuple[int, str, Table | Paragraph] | None = None
    match block:
        case CT_P():
            content_item = (i, "paragraph", Paragraph(block, doc))
        case CT_Tbl():
            content_item = (i, "table", Table(block, doc))
        case _:
            pass

    if content_item is not None:
        content_items.append(content_item)


def _normalize_docx_zip(file_content: bytes) -> bytes:
    r"""Some DOCX files (e.g. exported by Evernote on Windows) are malformed:
    ZIP entry names use backslash (\\) as path separator instead of the forward
    slash (/) required by both the ZIP spec and OOXML.  On Linux/Mac the entry
    "word\\document.xml" is never found when python-docx looks for
    "word/document.xml", which triggers a KeyError about a missing relationship.

    This function rewrites the ZIP in-memory, normalizing all entry names to
    use forward slashes without touching any actual document content.

    Returns:
        Normalized DOCX bytes, or the original bytes if the payload is not a ZIP.

    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_content), "r") as zin:
            out_buf = io.BytesIO()
            with zipfile.ZipFile(
                out_buf,
                "w",
                compression=zipfile.ZIP_DEFLATED,
            ) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    # Normalize backslash path separators to forward slash
                    item.filename = item.filename.replace("\\", "/")
                    zout.writestr(item, data)
            return out_buf.getvalue()
    except zipfile.BadZipFile:
        # Not a valid zip — return as-is and let python-docx report the real error
        return file_content


def _extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX file.
    For now support only paragraph and table add more if needed

    Returns:
        Extracted plain text content from the DOCX file.

    Raises:
        TextExtractionError: If the DOCX cannot be parsed after normalization and retry.

    """
    try:
        doc = _load_docx_document(file_content)
        text: list[str] = []
        for item_type, item in _iter_docx_content_items(doc):
            extracted_text = _extract_docx_item_text(item_type, item)
            if extracted_text is not None:
                text.append(extracted_text)
        return "\n".join(text)

    except Exception as e:
        logger.exception("Failed to extract text from DOCX")
        msg = f"Failed to extract text from DOCX: {e!s}"
        raise TextExtractionError(msg) from e


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a DOCX payload."""
    return _extract_text_from_docx(file_content)


def _load_docx_document(file_content: bytes) -> Document:
    try:
        return docx.Document(io.BytesIO(file_content))
    except Exception as e:
        logger.warning(
            "Failed to parse DOCX, attempting to normalize ZIP entry paths: %s",
            e,
            exc_info=e,
        )
        normalized_file_content = _normalize_docx_zip(file_content)
        return docx.Document(io.BytesIO(normalized_file_content))


def _iter_docx_content_items(doc: Document) -> list[tuple[str, Table | Paragraph]]:
    content_items: list[tuple[int, str, Table | Paragraph]] = []
    for index, part in enumerate(doc.element.body):
        parser_docx_part(part, doc, content_items, index)
    return [(item_type, item) for _, item_type, item in content_items]


def _extract_docx_item_text(item_type: str, item: Table | Paragraph) -> str | None:
    if item_type == "paragraph" and isinstance(item, Paragraph):
        return item.text
    if item_type == "table" and isinstance(item, Table):
        return _extract_docx_table_text(item)
    return None


def _extract_docx_table_text(item: Table) -> str | None:
    try:
        if not _docx_table_has_content(item):
            return None
        return _build_docx_markdown_table(item)
    except Exception as e:
        logger.warning("Failed to extract table from DOC: %s", e, exc_info=e)
        return None


def _docx_table_has_content(item: Table) -> bool:
    return any(any(cell.text.strip() for cell in row.cells) for row in item.rows)


def _build_docx_markdown_table(item: Table) -> str:
    cell_texts = [cell.text.replace("\n", "<br>") for cell in item.rows[0].cells]
    markdown_table = f"| {' | '.join(cell_texts)} |\n"
    markdown_table += f"| {' | '.join(['---'] * len(item.rows[0].cells))} |\n"
    for row in item.rows[1:]:
        row_cells = [cell.text.replace("\n", "<br>") for cell in row.cells]
        markdown_table += "| " + " | ".join(row_cells) + " |\n"
    return markdown_table


def _download_file_content(http_client: HttpClientProtocol, file: File) -> bytes:
    """Download the content of a file based on its transfer method."""
    if (
        file.transfer_method == FileTransferMethod.REMOTE_URL
        and file.remote_url is None
    ):
        msg = "Missing URL for remote file"
        raise FileDownloadError(msg)

    try:
        if (
            file.transfer_method == FileTransferMethod.REMOTE_URL
            and file.remote_url is not None
        ):
            response = http_client.get(file.remote_url)
            response.raise_for_status()
            return response.content
        return file_manager.download(file)
    except Exception as e:
        msg = f"Error downloading file: {e!s}"
        raise FileDownloadError(msg) from e


def download_file_content(http_client: HttpClientProtocol, file: File) -> bytes:
    """Download file content using the document extractor rules."""
    return _download_file_content(http_client, file)


def _extract_text_from_file(
    http_client: HttpClientProtocol,
    file: File,
    *,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    file_content = _download_file_content(http_client, file)
    if file.extension:
        extracted_text = _extract_text_by_file_extension(
            file_content=file_content,
            file_extension=file.extension,
            unstructured_api_config=unstructured_api_config,
        )
    elif file.mime_type:
        extracted_text = _extract_text_by_mime_type(
            file_content=file_content,
            mime_type=file.mime_type,
            unstructured_api_config=unstructured_api_config,
        )
    else:
        msg = "Unable to determine file type: MIME type or file extension is missing"
        raise UnsupportedFileTypeError(msg)
    return extracted_text


def _extract_text_from_csv(file_content: bytes) -> str:
    try:
        # Detect encoding using charset_normalizer
        result = charset_normalizer.from_bytes(file_content).best()
        encoding = result.encoding if result else "utf-8"

        # Fallback to utf-8 if detection fails
        if not encoding:
            encoding = "utf-8"

        try:
            csv_file = io.StringIO(file_content.decode(encoding, errors="ignore"))
        except (UnicodeDecodeError, LookupError):
            # If decoding fails, try with utf-8 as last resort
            csv_file = io.StringIO(file_content.decode("utf-8", errors="ignore"))

        csv_reader = csv.reader(csv_file)
        rows = list(csv_reader)

        if not rows:
            return ""

        # Combine multi-line text in the header row
        header_row = [cell.replace("\n", " ").replace("\r", "") for cell in rows[0]]

        # Create Markdown table
        markdown_table = "| " + " | ".join(header_row) + " |\n"
        markdown_table += (
            "| " + " | ".join(["-" * len(col) for col in rows[0]]) + " |\n"
        )

        # Process each data row and combine multi-line text in each cell
        for row in rows[1:]:
            processed_row = [cell.replace("\n", " ").replace("\r", "") for cell in row]
            markdown_table += "| " + " | ".join(processed_row) + " |\n"

    except Exception as e:
        msg = f"Failed to extract text from CSV: {e!s}"
        raise TextExtractionError(msg) from e
    else:
        return markdown_table


def _extract_text_from_excel(file_content: bytes) -> str:
    """Extract text from an Excel file using pandas."""

    def _construct_markdown_table(df: pd.DataFrame) -> str:
        """Manually construct a Markdown table from a DataFrame."""
        # Construct the header row
        header_row = "| " + " | ".join(df.columns) + " |"

        # Construct the separator row
        separator_row = "| " + " | ".join(["-" * len(col) for col in df.columns]) + " |"

        # Construct the data rows
        data_rows = []
        for _, row in df.iterrows():
            data_row = "| " + " | ".join(map(str, row)) + " |"
            data_rows.append(data_row)

        # Combine all rows into a single string
        return "\n".join([header_row, separator_row, *data_rows])

    try:
        excel_file = pd.ExcelFile(io.BytesIO(file_content))
        markdown_table = ""
        for sheet_name in excel_file.sheet_names:
            try:
                df = excel_file.parse(sheet_name=sheet_name)
                if not isinstance(df, pd.DataFrame):
                    continue
                df = df.dropna(how="all")

                # Combine multi-line text in each cell into a single line
                df = df.map(
                    lambda x: (
                        " ".join(str(x).splitlines()) if isinstance(x, str) else x
                    ),
                )

                # Combine multi-line text in column names into a single line
                df.columns = pd.Index([
                    " ".join(str(col).splitlines()) for col in df.columns
                ])

                # Manually construct the Markdown table
                markdown_table += _construct_markdown_table(df) + "\n\n"
            except (TypeError, ValueError):
                continue
    except Exception as e:
        msg = f"Failed to extract text from Excel file: {e!s}"
        raise TextExtractionError(msg) from e
    else:
        return markdown_table


def _extract_text_from_ppt(
    file_content: bytes,
    *,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    try:
        if unstructured_api_config.api_url:
            from unstructured.partition.api import partition_via_api  # noqa: PLC0415

            elements = _partition_file_via_unstructured_api(
                partition_via_api,
                file_content,
                suffix=".ppt",
                unstructured_api_config=unstructured_api_config,
            )
        else:
            from unstructured.partition.ppt import partition_ppt  # noqa: PLC0415

            with io.BytesIO(file_content) as file:
                elements = partition_ppt(file=file)
        return "\n".join([getattr(element, "text", "") for element in elements])

    except Exception as e:
        msg = f"Failed to extract text from PPTX: {e!s}"
        raise TextExtractionError(msg) from e


def _extract_text_from_pptx(
    file_content: bytes,
    *,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    try:
        if unstructured_api_config.api_url:
            from unstructured.partition.api import partition_via_api  # noqa: PLC0415

            elements = _partition_file_via_unstructured_api(
                partition_via_api,
                file_content,
                suffix=".pptx",
                unstructured_api_config=unstructured_api_config,
            )
        else:
            from unstructured.partition.pptx import partition_pptx  # noqa: PLC0415

            with io.BytesIO(file_content) as file:
                elements = partition_pptx(file=file)
        return "\n".join([getattr(element, "text", "") for element in elements])
    except Exception as e:
        msg = f"Failed to extract text from PPTX: {e!s}"
        raise TextExtractionError(msg) from e


def _extract_text_from_epub(
    file_content: bytes,
    *,
    unstructured_api_config: UnstructuredApiConfig,
) -> str:
    try:
        if unstructured_api_config.api_url:
            from unstructured.partition.api import partition_via_api  # noqa: PLC0415

            elements = _partition_file_via_unstructured_api(
                partition_via_api,
                file_content,
                suffix=".epub",
                unstructured_api_config=unstructured_api_config,
            )
        else:
            pypandoc.download_pandoc()
            from unstructured.partition.epub import partition_epub  # noqa: PLC0415

            with io.BytesIO(file_content) as file:
                elements = partition_epub(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        msg = f"Failed to extract text from EPUB: {e!s}"
        raise TextExtractionError(msg) from e


def _extract_text_from_eml(file_content: bytes) -> str:
    try:
        from unstructured.partition.email import partition_email  # noqa: PLC0415

        with io.BytesIO(file_content) as file:
            elements = partition_email(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        msg = f"Failed to extract text from EML: {e!s}"
        raise TextExtractionError(msg) from e


def _extract_text_from_msg(file_content: bytes) -> str:
    try:
        from unstructured.partition.msg import partition_msg  # noqa: PLC0415

        with io.BytesIO(file_content) as file:
            elements = partition_msg(file=file)
        return "\n".join([str(element) for element in elements])
    except Exception as e:
        msg = f"Failed to extract text from MSG: {e!s}"
        raise TextExtractionError(msg) from e


def _extract_text_from_vtt(vtt_bytes: bytes) -> str:
    text = _extract_text_from_plain_text(vtt_bytes)

    # remove bom
    text = text.lstrip("\ufeff")

    raw_results = [
        (caption.voice, caption.text) for caption in webvtt.from_string(text)
    ]

    # Merge consecutive utterances by the same speaker
    merged_results = []
    if raw_results:
        current_speaker, current_text = raw_results[0]

        for i in range(1, len(raw_results)):
            spk, txt = raw_results[i]
            if spk is None:
                merged_results.append((None, current_text))
                continue

            if spk == current_speaker:
                # If it is the same speaker, merge the utterances (joined by space)
                current_text += " " + txt
            else:
                # If the speaker changes, register the utterance so far and move on
                merged_results.append((current_speaker, current_text))
                current_speaker, current_text = spk, txt

        # Add the last element
        merged_results.append((current_speaker, current_text))
    else:
        merged_results = raw_results

    # Return the result in the specified format: Speaker "text" style
    formatted = [f'{spk or ""} "{txt}"' for spk, txt in merged_results]
    return "\n".join(formatted)


def _extract_text_from_properties(file_content: bytes) -> str:
    try:
        text = _extract_text_from_plain_text(file_content)
        lines = text.splitlines()
        result = []
        for line in lines:
            stripped_line = line.strip()
            # Preserve comments and empty lines
            if not stripped_line or stripped_line.startswith(("#", "!")):
                result.append(stripped_line)
                continue

            if "=" in stripped_line:
                key, value = stripped_line.split("=", 1)
            elif ":" in stripped_line:
                key, value = stripped_line.split(":", 1)
            else:
                key, value = stripped_line, ""

            result.append(f"{key.strip()}: {value.strip()}")

        return "\n".join(result)
    except Exception as e:
        msg = f"Failed to extract text from properties file: {e!s}"
        raise TextExtractionError(msg) from e


def _build_text_extractor_registry() -> _TextExtractorRegistry:
    return _TextExtractorRegistry((
        _ExtractorRegistration(
            name="plain_text",
            extractor=_extract_text_from_plain_text,
            mime_types=_MIME_PLAIN_TEXT_TYPES,
            file_extensions=_EXTENSION_PLAIN_TEXT_TYPES,
        ),
        _ExtractorRegistration(
            name="json",
            extractor=_extract_text_from_json,
            mime_types=frozenset({"application/json"}),
            file_extensions=frozenset({".json"}),
        ),
        _ExtractorRegistration(
            name="yaml",
            extractor=_extract_text_from_yaml,
            mime_types=frozenset({"application/x-yaml", "text/yaml"}),
            file_extensions=frozenset({".yaml", ".yml"}),
        ),
        _ExtractorRegistration(
            name="pdf",
            extractor=_extract_text_from_pdf,
            mime_types=frozenset({"application/pdf"}),
            file_extensions=frozenset({".pdf"}),
        ),
        _ExtractorRegistration(
            name="doc",
            extractor=_extract_text_from_doc,
            mime_types=frozenset({"application/msword"}),
            file_extensions=frozenset({".doc"}),
            requires_unstructured_config=True,
        ),
        _ExtractorRegistration(
            name="docx",
            extractor=_extract_text_from_docx,
            mime_types=frozenset({
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }),
            file_extensions=frozenset({".docx"}),
        ),
        _ExtractorRegistration(
            name="csv",
            extractor=_extract_text_from_csv,
            mime_types=frozenset({"text/csv"}),
            file_extensions=frozenset({".csv"}),
        ),
        _ExtractorRegistration(
            name="excel",
            extractor=_extract_text_from_excel,
            mime_types=frozenset({
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            }),
            file_extensions=frozenset({".xls", ".xlsx"}),
        ),
        _ExtractorRegistration(
            name="ppt",
            extractor=_extract_text_from_ppt,
            mime_types=frozenset({"application/vnd.ms-powerpoint"}),
            file_extensions=frozenset({".ppt"}),
            requires_unstructured_config=True,
        ),
        _ExtractorRegistration(
            name="pptx",
            extractor=_extract_text_from_pptx,
            mime_types=frozenset({
                "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            }),
            file_extensions=frozenset({".pptx"}),
            requires_unstructured_config=True,
        ),
        _ExtractorRegistration(
            name="epub",
            extractor=_extract_text_from_epub,
            mime_types=frozenset({"application/epub+zip"}),
            file_extensions=frozenset({".epub"}),
            requires_unstructured_config=True,
        ),
        _ExtractorRegistration(
            name="eml",
            extractor=_extract_text_from_eml,
            mime_types=frozenset({"message/rfc822"}),
            file_extensions=frozenset({".eml"}),
        ),
        _ExtractorRegistration(
            name="msg",
            extractor=_extract_text_from_msg,
            mime_types=frozenset({"application/vnd.ms-outlook"}),
            file_extensions=frozenset({".msg"}),
        ),
        _ExtractorRegistration(
            name="vtt",
            extractor=_extract_text_from_vtt,
            mime_types=frozenset({"text/vtt"}),
        ),
        _ExtractorRegistration(
            name="properties",
            extractor=_extract_text_from_properties,
            mime_types=frozenset({"text/properties"}),
            file_extensions=frozenset({".properties"}),
        ),
    ))


_TEXT_EXTRACTOR_REGISTRY = _build_text_extractor_registry()
